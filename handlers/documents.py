import asyncio
import os
import uuid
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from aiogram.types import Message, FSInputFile
from aiogram.fsm.context import FSMContext

from aiogram.utils.keyboard import InlineKeyboardBuilder
from config import GENERATED_DIR, SUPPORT_CONTACT, AI_API_KEY, PAYMENT_ENABLED
from db import save_case

TODAY = lambda: date.today().strftime("%d.%m.%Y")


def _set_doc_margins(doc: Document):
    """Поля по стандарту делопроизводства РК: лево 3см, право 1.5см, верх/низ 2см."""
    for section in doc.sections:
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Убираем дефолтный отступ после абзаца во всём документе
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.line_spacing = Pt(14)  # 1.15 для 12pt
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


async def _clean_description(raw_text: str) -> str:
    """
    Переписывает сырой текст пользователя в юридически грамотный
    язык для официального заявления. Если AI недоступен — возвращает
    исходный текст.
    """
    if not AI_API_KEY or AI_API_KEY in ("вставьте_ключ_groq_сюда", ""):
        return raw_text  # AI не настроен — используем как есть

    try:
        from ai_service import get_legal_assessment
        from openai import AsyncOpenAI
        from config import AI_BASE_URL, AI_MODEL

        client = AsyncOpenAI(api_key=AI_API_KEY, base_url=AI_BASE_URL)

        async def _call():
            return await client.chat.completions.create(
                model=AI_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Ты помощник юриста в Казахстане. Перепиши описание ситуации "
                            "буллинга для официального заявления директору школы. "
                            "Требования: официальный деловой стиль, без эмоций, "
                            "конкретные факты, исправь опечатки и грамматику, "
                            "убери повторы, длина не более 5-7 предложений. "
                            "Верни ТОЛЬКО переписанный текст, без пояснений."
                        ),
                    },
                    {"role": "user", "content": raw_text},
                ],
                temperature=0.2,
                max_tokens=400,
            )

        response = await asyncio.wait_for(_call(), timeout=15.0)
        cleaned = response.choices[0].message.content.strip()
        return cleaned if cleaned else raw_text
    except (asyncio.TimeoutError, Exception):
        return raw_text  # если что-то пошло не так — используем оригинал


# ── Helpers ───────────────────────────────────────────────────────────────

def _set_font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _set_spacing(paragraph, space_after_pt=0):
    """Одинарный межстрочный интервал, нулевые отступы, выравнивание по ширине."""
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(14)        # ~1.15 для 12pt — читаемо, не разреженно
    pf.space_before = Pt(0)
    pf.space_after  = Pt(space_after_pt)
    pf.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY


def _heading(doc: Document, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    _set_font(run, bold=True, size=13)


def _body(doc: Document, text: str, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        _set_spacing(p)
    run = p.add_run(text)
    _set_font(run, bold=bold, italic=italic)
    return p


def _save(doc: Document, filename: str) -> str:
    os.makedirs(GENERATED_DIR, exist_ok=True)
    path = os.path.join(GENERATED_DIR, filename)
    doc.save(path)
    return path


# ── Document 1: Заявление директору школы ─────────────────────────────────

def make_doc1_school_application(data: dict) -> str:
    doc = Document()
    _set_doc_margins(doc)

    school = data.get('school_name') or '___________'
    city   = data.get('city') or '___________'
    child  = data.get('child_name') or '___________'
    cls    = data.get('child_class') or '___'
    appl   = data.get('applicant_name') or '___________'

    # Шапка (кому/от кого) — правый угол
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run(
        f"Директору {school}\n"
        f"г. {city}\n\n"
        f"От: {appl},\n"
        f"законного представителя обучающегося\n"
        f"{child}, {cls} класс"
    )
    _set_font(r)

    doc.add_paragraph()
    _heading(doc, "Заявление\nо регистрации факта травли (буллинга)")
    doc.add_paragraph()

    _body(doc,
        f"Я, {appl}, законный представитель обучающегося "
        f"{child} ({cls} класс), обращаюсь к Вам "
        f"в соответствии с Правилами профилактики травли (буллинга) ребёнка, "
        f"утверждёнными Приказом Министра просвещения РК от 21.12.2022 № 506 "
        f"(далее — Правила № 506)."
    )

    doc.add_paragraph()
    _body(doc, "Обстоятельства:", bold=True)

    desc = data.get('incident_description', '').strip()
    _body(doc,
        f"В отношении моего ребёнка {child} имели место следующие действия: {desc}. "
        f"Приложения: при наличии — скриншоты, фото, медицинские справки, иные документы."
    )

    doc.add_paragraph()
    _body(doc, "ПРОШУ:", bold=True)

    reqs = [
        "Зарегистрировать настоящее обращение в журнале учёта информации о фактах травли (буллинга) "
        "в соответствии с Правилами № 506 и сообщить входящий номер и дату регистрации.",
        "Незамедлительно начать процедуру реагирования: сбор первичной информации, "
        "попытку медиативного урегулирования, и при необходимости — вынесение вопроса "
        "на рассмотрение Совета профилактики.",
        "Обеспечить письменное информирование о принятых мерах и результатах рассмотрения "
        "в сроки, предусмотренные Правилами № 506 и АППК РК.",
    ]
    for i, req in enumerate(reqs, 1):
        _body(doc, f"{i}. {req}")

    doc.add_paragraph()
    _body(doc,
        "Обращаю внимание: в соответствии с АППК РК (ст. 64) обращение подлежит обязательному "
        "приёму и регистрации. Отказ в приёме не допускается.",
        italic=True
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"«___» __________ 2026 г.            Подпись: _______________  / {appl} /")
    _set_font(r)

    filename = f"1_zayavlenie_shkola_{uuid.uuid4().hex[:8]}.docx"
    return _save(doc, filename)


# ── Document 2: Требование письменного ответа (АППК) ─────────────────────

def make_doc2_written_response(data: dict) -> str:
    doc = Document()
    _set_doc_margins(doc)

    school = data.get('school_name') or '___________'
    city   = data.get('city') or '___________'
    child  = data.get('child_name') or '___________'
    cls    = data.get('child_class') or '___'
    appl   = data.get('applicant_name') or '___________'

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run(
        f"Директору {school}\n"
        f"г. {city}\n\n"
        f"От: {appl},\n"
        f"законного представителя {child},\n"
        f"{cls} класс"
    )
    _set_font(r)

    doc.add_paragraph()
    _heading(doc, "Требование\nо предоставлении письменной информации о результатах\nрассмотрения обращения о травле (буллинге)")
    doc.add_paragraph()

    _body(doc,
        f"«___» __________ 2026 г. мной было подано письменное заявление о регистрации "
        f"факта травли в отношении {child} (вх. № _______). "
        f"На сегодняшний день письменный ответ о принятых мерах мной не получен / "
        f"получен ненадлежащим образом."
    )
    doc.add_paragraph()

    _body(doc,
        "В соответствии с Административным процедурно-процессуальным кодексом РК (АППК РК) "
        "и Правилами № 506, ТРЕБУЮ:",
        bold=True
    )

    reqs = [
        "Предоставить письменный ответ о результатах рассмотрения моего обращения, "
        "в том числе о решении Совета профилактики (признание/непризнание факта травли) "
        "и перечне принятых мер.",
        "Подтвердить факт и сроки передачи информации в местный исполнительный орган "
        "в сфере образования в соответствии с Правилами № 506.",
        "В случае непринятия мер — указать мотивированное обоснование.",
    ]
    for i, req in enumerate(reqs, 1):
        _body(doc, f"{i}. {req}")

    doc.add_paragraph()
    _body(doc,
        "Предупреждаю: в случае отсутствия надлежащего ответа буду вынужден(а) обратиться "
        "в местный исполнительный орган в сфере образования и органы прокуратуры.",
        italic=True
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"«___» __________ 2026 г.            Подпись: _______________  / {appl} /")
    _set_font(r)

    filename = f"2_trebovanie_otvet_{uuid.uuid4().hex[:8]}.docx"
    return _save(doc, filename)


# ── Document 3: Памятка по кибербуллингу ─────────────────────────────────

def make_doc3_cyberbullying_memo(data: dict) -> str:
    doc = Document()
    _set_doc_margins(doc)

    appl  = data.get('applicant_name') or '___________'
    child = data.get('child_name') or '___________'

    _heading(doc, "Памятка по фиксации кибербуллинга")
    _body(doc, f"Составлена для: {appl}, законного представителя {child}", bold=True)
    doc.add_paragraph()

    _body(doc, "1. ЧТО ФИКСИРОВАТЬ", bold=True)
    items1 = [
        "Скриншоты сообщений — с видимым именем/ником отправителя, датой и временем.",
        "Ссылки на публикации, группы, аккаунты, где происходит преследование.",
        "Видеозаписи (экран телефона/компьютера), если публикации могут быть удалены.",
        "Переписку, голосовые сообщения, фото/видео угрожающего характера.",
        "Сведения о платформе (ВКонтакте, Instagram, Telegram, TikTok и т.д.).",
    ]
    for item in items1:
        _body(doc, f"• {item}")

    doc.add_paragraph()
    _body(doc, "2. КАК ФИКСИРОВАТЬ ПРАВИЛЬНО", bold=True)
    items2 = [
        "Делайте скриншоты ДО того, как обидчик удалит сообщения.",
        "НЕ редактируйте скриншоты — нужны оригиналы (без обрезки, стикеров, пометок).",
        "Сохраняйте файлы с оригинальными датами создания (не делайте повторные скриншоты).",
        "Ведите хронологический список: дата — платформа — суть — ссылка/файл.",
        "Скопируйте ссылки на профили обидчиков (URL страниц).",
    ]
    for item in items2:
        _body(doc, f"• {item}")

    doc.add_paragraph()
    _body(doc, "3. КУДА СООБЩАТЬ О КИБЕРБУЛЛИНГЕ", bold=True)
    items3 = [
        "В школу: письменное заявление директору (Правила № 506 распространяются на кибербуллинг).",
        "В администрацию платформы: функция «Пожаловаться» / Report.",
        "В полицию (102): если есть угрозы жизни, вымогательство, «сексторшн».",
        "Уполномоченному по правам ребёнка РК — при системном бездействии органов.",
    ]
    for item in items3:
        _body(doc, f"• {item}")

    doc.add_paragraph()
    _body(doc, "4. ПРАВОВАЯ ОСНОВА", bold=True)
    _body(doc,
        "Кибербуллинг прямо включён в определение травли (буллинга) по Правилам № 506 "
        "(«через публичность/медиа/сети/онлайн-платформы»). Школа обязана реагировать "
        "вне зависимости от того, произошло ли это в стенах школы или в интернете."
    )

    doc.add_paragraph()
    _body(doc, "5. ВАЖНО ПОМНИТЬ", bold=True)
    _body(doc,
        "НЕ требуйте от ребёнка «не реагировать» и «терпеть» — это не решает проблему. "
        "НЕ удаляйте переписку сами — это уничтожает доказательства. "
        "Поддержите ребёнка: объясните, что он не виноват в произошедшем."
    )

    doc.add_paragraph()
    _body(doc, f"Дата составления: {TODAY()}")

    filename = f"3_pamyatka_kiberbulling_{uuid.uuid4().hex[:8]}.docx"
    return _save(doc, filename)


# ── Document 4: Жалоба в акимат / отдел образования ─────────────────────

def make_doc4_akimat_complaint(data: dict) -> str:
    doc = Document()
    _set_doc_margins(doc)

    school = data.get('school_name') or '___________'
    city   = data.get('city') or '___________'
    child  = data.get('child_name') or '___________'
    cls    = data.get('child_class') or '___'
    appl   = data.get('applicant_name') or '___________'
    desc   = data.get('incident_description', '').strip()

    # Шапка — правый угол
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = header.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = Pt(14)
    r = header.add_run(
        f"Руководителю районного (городского)\n"
        f"отдела образования г. {city}\n\n"
        f"От: {appl},\n"
        f"законного представителя обучающегося\n"
        f"{child}, {cls} класс,\n"
        f"{school}"
    )
    _set_font(r)

    doc.add_paragraph()
    _heading(doc, "Жалоба\nна бездействие администрации общеобразовательной организации")
    doc.add_paragraph()

    _body(doc,
        f"Я, {appl}, законный представитель обучающегося {child} ({cls} класс, {school}), "
        f"обращаюсь в соответствии с Правилами профилактики и реагирования на проявления травли (буллинга), "
        f"утверждёнными Приказом Министра просвещения РК от 21.12.2022 № 506 (далее — Правила № 506), "
        f"а также на основании Административного процедурно-процессуального кодекса РК (АППК РК)."
    )

    doc.add_paragraph()
    _body(doc, "Обстоятельства:", bold=True)
    _body(doc,
        f"В отношении моего ребёнка {child} имели место следующие действия: {desc}. "
        f"«___» __________ 2026 г. мной было подано письменное заявление директору {school} "
        f"о регистрации факта травли (буллинга) в соответствии с Правилами № 506."
    )

    doc.add_paragraph()
    _body(doc, "Суть жалобы:", bold=True)
    _body(doc,
        f"По истечении установленного срока письменный ответ от администрации {school} "
        f"мной получен не был / полученный ответ не содержит сведений о конкретных мерах "
        f"реагирования, предусмотренных Правилами № 506. Таким образом, имеет место бездействие "
        f"администрации школы, нарушающее права моего ребёнка на защиту от травли."
    )

    doc.add_paragraph()
    _body(doc, "ПРОШУ:", bold=True)

    reqs = [
        f"Провести проверку соблюдения администрацией {school} требований Правил № 506 "
        f"в части регистрации и рассмотрения моего обращения.",
        "Обязать администрацию школы провести надлежащую процедуру реагирования на факт буллинга: "
        "сбор первичной информации, медиацию или рассмотрение на Совете профилактики.",
        "Предоставить мне письменный ответ о результатах проверки и принятых мерах "
        "в сроки, предусмотренные АППК РК.",
        "При подтверждении бездействия — принять меры административного воздействия "
        "в отношении должностных лиц школы.",
    ]
    for i, req in enumerate(reqs, 1):
        _body(doc, f"{i}. {req}")

    doc.add_paragraph()
    _body(doc,
        "Приложения: копия заявления директору школы с отметкой о принятии (при наличии); "
        "иные имеющиеся доказательства.",
        italic=True
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"«___» __________ 2026 г.            Подпись: _______________  / {appl} /")
    _set_font(r)

    filename = f"4_zhaloba_akimat_{uuid.uuid4().hex[:8]}.docx"
    return _save(doc, filename)


# ── Main: generate all docs and send ─────────────────────────────────────

async def generate_and_send(message: Message, state: FSMContext):
    from db import create_reminders
    data = await state.get_data()
    user_id = message.chat.id

    await message.answer("⏳ Обрабатываю и генерирую документы...")

    # Переписываем описание через LLM (юридический стиль)
    raw_desc = data.get("incident_description", "")
    if raw_desc:
        data["incident_description"] = await _clean_description(raw_desc)

    # Save to DB
    case_id = await save_case(user_id, data)

    # Generate all 4 documents
    paths = [
        make_doc1_school_application(data),
        make_doc2_written_response(data),
        make_doc3_cyberbullying_memo(data),
        make_doc4_akimat_complaint(data),
    ]

    # ── Документ 1 — всегда бесплатно ────────────────────────────────────
    await message.answer("📄 <b>Документ 1: Заявление директору школы</b> — бесплатно", parse_mode="HTML")
    await message.answer_document(FSInputFile(paths[0]))

    # ── Документы 2–4 — freemium ─────────────────────────────────────────
    if PAYMENT_ENABLED:
        kb = InlineKeyboardBuilder()
        kb.button(
            text="📦 Базовый — 2 990 тг (Документы 1–3)",
            callback_data=f"pay_b_{case_id}",
        )
        kb.button(
            text="⭐ Полный — 4 990 тг (Документы 1–4 + напоминания)",
            callback_data=f"pay_f_{case_id}",
        )
        kb.adjust(1)
        await message.answer(
            "📦 <b>Выберите пакет документов:</b>\n\n"
            "🔹 <b>Базовый — 2 990 тг</b>\n"
            "   Документы 1–3: заявление + требование ответа + памятка\n\n"
            "⭐ <b>Полный — 4 990 тг</b>\n"
            "   Документы 1–4 + автонапоминания на 5-й и 10-й день\n"
            "   (бот напомнит если школа не ответила и выдаст жалобу в акимат)",
            reply_markup=kb.as_markup(),
            parse_mode="HTML",
        )
    else:
        # Режим разработки — выдаём все документы
        doc_labels = [
            "📄 <b>Документ 2: Требование письменного ответа</b>",
            "📄 <b>Документ 3: Памятка по фиксации кибербуллинга</b>",
            "📄 <b>Документ 4: Жалоба в акимат / отдел образования</b>",
        ]
        for path, label in zip(paths[1:], doc_labels):
            await message.answer(label, parse_mode="HTML")
            await message.answer_document(FSInputFile(path))

        # В тестовом режиме — создаём напоминания автоматически
        await create_reminders(user_id, case_id)

    # ── Советы по уровню ситуации ────────────────────────────────────────
    level = data.get("triage_level", "GREEN")
    level_tip = {
        "GREEN": (
            "🟢 <b>Следующие шаги:</b>\n"
            "1. Распечатайте Документ 1 и подайте директору школы — требуйте входящий номер\n"
            "2. Если в течение 5 рабочих дней нет ответа — используйте Документ 2\n"
            "3. Если школа не реагирует — обращайтесь в отдел образования акимата (Документ 4)"
        ),
        "AMBER": (
            "🟡 <b>Следующие шаги:</b>\n"
            "1. Подайте Документ 1 в школу — и параллельно обратитесь в полицию (102)\n"
            "2. Сохраните все доказательства (Документ 3 — инструкция)\n"
            "3. При вымогательстве/насилии — медицинская справка обязательна\n"
            "4. Возраст обидчика {age}: {age_tip}"
        ),
        "RED": (
            "🔴 <b>Экстренный уровень:</b>\n"
            "1. Сначала — безопасность ребёнка и медицинская помощь\n"
            "2. Полиция (102) и скорая (103) — ДО школы\n"
            "3. Документы из этого пакета используйте после урегулирования острой ситуации"
        ),
    }.get(level, "")

    if level == "AMBER":
        age = data.get("bully_age_group", "")
        if "14–15" in age:
            age_tip = "уголовная ответственность возможна по ст.15 УК РК"
        elif "до 14" in age.lower():
            age_tip = "уголовная ответственность не наступает — трек через школу и отдел образования акимата"
        else:
            age_tip = "уточните возраст для выбора правового трека"
        level_tip = level_tip.format(age=age, age_tip=age_tip)

    if level_tip:
        await message.answer(level_tip, parse_mode="HTML")

    # ── Дисклеймер ───────────────────────────────────────────────────────
    await message.answer(
        "⚠️ <b>Дисклеймер:</b> Документы сформированы на основе Правил № 506 МП РК. "
        "Помогаем родителям защитить детей — разбираемся в законах и готовим документы. "
        "Для сложных ситуаций рядом есть юристы ADDASTRA.\n\n"
        "Введите /start чтобы начать новое обращение.",
        parse_mode="HTML",
    )

    # ── ADDASTRA: предложение консультации ───────────────────────────────
    kb2 = InlineKeyboardBuilder()
    kb2.button(text="💼 Получить консультацию юриста", callback_data="consult_addastra")
    kb2.adjust(1)
    await message.answer(
        "👨‍⚖️ <b>Нужна помощь юриста ADDASTRA?</b>\n\n"
        "Если ситуация сложная и вы хотите, чтобы юрист лично сопроводил дело — "
        "оставьте заявку. Мы свяжемся с вами в течение рабочего дня.",
        reply_markup=kb2.as_markup(),
        parse_mode="HTML",
    )

    await state.clear()

    # Clean up generated files
    for path in paths:
        try:
            os.remove(path)
        except OSError:
            pass
